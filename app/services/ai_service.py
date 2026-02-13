"""Сервис для работы с Google Generative AI (Gemini)."""
import io
from pathlib import Path
from typing import List, Optional

from google import genai
from google.genai import types
from google.api_core.exceptions import ResourceExhausted
import numpy as np
from sqlalchemy.ext.asyncio import AsyncSession

try:
    import fitz  # PyMuPDF для чтения PDF (импортируется как fitz)
except ImportError:
    fitz = None  # type: ignore

try:
    from pydub import AudioSegment
    PYDUB_AVAILABLE = True
    # Примечание: pydub требует ffmpeg для конвертации аудио форматов
    # Если ffmpeg не установлен, конвертация не будет работать, но это не критично
except ImportError:
    PYDUB_AVAILABLE = False
    AudioSegment = None  # type: ignore

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None  # type: ignore

from app.core.config import settings
from app.services.vector_store import VectorStore
from app.services.chat_history import (
    get_recent_messages,
    save_message,
    format_history_for_prompt,
)
from app.utils.logger import logger

# Сообщение при превышении квоты
QUOTA_EXCEEDED_MESSAGE = "⚠️ Слишком много вопросов! Мозгу нужно отдохнуть 15 секунд. Пожалуйста, повтори запрос чуть позже."

# Инициализация клиента Google GenAI (если ключ задан)
gemini_client: Optional[genai.Client] = None
if settings.gemini_api_key:
    try:
        gemini_client = genai.Client(api_key=settings.gemini_api_key)
        
        # Вывод всех доступных моделей для отладки
        available_models = [m.name for m in gemini_client.models.list()]
        print(f"[DEBUG] Available Gemini models: {available_models}")
        logger.info(f"[DEBUG] Available Gemini models: {available_models}")
    except Exception as e:
        logger.warning(f"[DEBUG] Failed to initialize Gemini client: {e}")
        gemini_client = None


class GeminiService:
    """Сервис для работы с Google Generative AI (Gemini)."""
    
    # Векторные хранилища по отделам (multitenancy)
    # Формат: {'delivery/courier': VectorStore, 'sorting': VectorStore, ...}
    _vector_stores: dict[str, VectorStore] = {}
    
    # Старое хранилище для обратной совместимости (deprecated)
    _vector_store: VectorStore | None = None
    
    @staticmethod
    def rebuild_index_for_department(department: str) -> None:
        """
        Точечное обновление индекса конкретного отдела (оптимизация).
        
        Args:
            department: Название отдела (например, 'sorting', 'manager', 'delivery/courier')
        """
        try:
            logger.info(f"[RAG] 🎯 Точечное обновление индекса для отдела: {department}")
            
            from app.core.models import Department as DepartmentEnum
            import fitz
            
            knowledge_path = Path("data/knowledge")
            text_extensions = {".txt", ".md", ".rst"}
            pdf_extensions = {".pdf"}
            docx_extensions = {".docx"}
            
            # Сначала загружаем common файлы
            common_path = knowledge_path / "common"
            common_chunks: List[str] = []
            common_metadata: List[dict] = []
            
            if common_path.exists() and common_path.is_dir():
                for file_path in common_path.iterdir():
                    if not file_path.is_file():
                        continue
                    
                    file_ext = file_path.suffix.lower()
                    content = ""
                    
                    try:
                        if file_ext in text_extensions:
                            content = file_path.read_text(encoding="utf-8")
                        elif file_ext in pdf_extensions:
                            if fitz is None:
                                continue
                            doc = fitz.open(file_path)
                            content = "\n".join([page.get_text() for page in doc])
                            doc.close()
                        elif file_ext in docx_extensions:
                            if not DOCX_AVAILABLE or Document is None:
                                continue
                            doc = Document(file_path)
                            content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                        else:
                            continue
                        
                        if content:
                            file_chunks = GeminiService._split_text_into_chunks(content, chunk_size=1000, overlap=200)
                            common_chunks.extend(file_chunks)
                            common_metadata.extend([{"filename": f"common/{file_path.name}"} for _ in file_chunks])
                    
                    except Exception as e:
                        logger.warning(f"[RAG] Failed to process common/{file_path.name}: {e}")
            
            # Теперь загружаем файлы конкретного отдела
            dept_chunks: List[str] = []
            dept_metadata: List[dict] = []
            
            dept_chunks.extend(common_chunks)
            dept_metadata.extend(common_metadata)
            
            dept_path = knowledge_path / department
            if dept_path.exists() and dept_path.is_dir():
                for file_path in dept_path.rglob("*"):
                    if not file_path.is_file():
                        continue
                    
                    file_ext = file_path.suffix.lower()
                    content = ""
                    
                    try:
                        if file_ext in text_extensions:
                            content = file_path.read_text(encoding="utf-8")
                        elif file_ext in pdf_extensions:
                            if fitz is None:
                                continue
                            doc = fitz.open(file_path)
                            content = "\n".join([page.get_text() for page in doc])
                            doc.close()
                        elif file_ext in docx_extensions:
                            if not DOCX_AVAILABLE or Document is None:
                                continue
                            doc = Document(file_path)
                            content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                        else:
                            continue
                        
                        if content:
                            file_chunks = GeminiService._split_text_into_chunks(content, chunk_size=1000, overlap=200)
                            dept_chunks.extend(file_chunks)
                            dept_metadata.extend([{"filename": f"{department}/{file_path.name}"} for _ in file_chunks])
                            logger.info(f"[RAG] Processed {department}/{file_path.name}: {len(file_chunks)} chunks")
                    
                    except Exception as e:
                        logger.warning(f"[RAG] Failed to process {department}/{file_path.name}: {e}")
            
            # Создаем/обновляем индекс отдела
            if dept_chunks:
                logger.info(f"[RAG] Department {department}: {len(dept_chunks)} chunks total")
                embeddings = GeminiService._generate_embeddings(dept_chunks)
                
                if len(embeddings) != len(dept_chunks):
                    dept_chunks = dept_chunks[:embeddings.shape[0]]
                    dept_metadata = dept_metadata[:embeddings.shape[0]]
                
                vector_store = VectorStore()
                vector_store.clear()
                vector_store.add_embeddings(embeddings, dept_chunks, dept_metadata)
                GeminiService._vector_stores[department] = vector_store
                logger.info(f"[RAG] ✅ Index updated for {department}: {len(dept_chunks)} chunks")
            else:
                logger.warning(f"[RAG] No chunks for department {department}")
        
        except Exception as e:
            logger.error(f"[RAG] Error rebuilding index for {department}: {e}", exc_info=True)
    
    @staticmethod
    def _create_department_indices() -> None:
        """
        Создает отдельные векторные индексы для каждого отдела (multitenancy).
        Каждый индекс включает: файлы отдела + файлы из common/.
        """
        try:
            logger.info("[RAG] Creating department-based vector indices...")
            
            from app.core.models import Department
            
            knowledge_path = Path("data/knowledge")
            if not knowledge_path.exists():
                logger.warning("[RAG] Knowledge base directory not found")
                GeminiService._vector_stores = {}
                return
            
            # Поддерживаемые форматы
            text_extensions = {".txt", ".md", ".rst"}
            pdf_extensions = {".pdf"}
            docx_extensions = {".docx"}
            
            # Получаем список всех отделов
            departments = [dept.value for dept in Department]
            
            # Сначала читаем common файлы (они будут добавлены во все индексы)
            common_path = knowledge_path / "common"
            common_chunks: List[str] = []
            common_metadata: List[dict] = []
            
            if common_path.exists() and common_path.is_dir():
                logger.info("[RAG] Loading common knowledge...")
                for file_path in common_path.iterdir():
                    if not file_path.is_file():
                        continue
                    
                    file_ext = file_path.suffix.lower()
                    content = ""
                    
                    try:
                        # Читаем текстовые файлы
                        if file_ext in text_extensions:
                            content = file_path.read_text(encoding="utf-8")
                        
                        # Читаем PDF файлы
                        elif file_ext in pdf_extensions:
                            if fitz is None:
                                logger.warning(f"[RAG] PyMuPDF not installed, skipping PDF: {file_path.name}")
                                continue
                            
                            doc = fitz.open(file_path)
                            text_parts: List[str] = []
                            for page in doc:
                                text_parts.append(page.get_text())
                            content = "\n".join(text_parts)
                            doc.close()
                        
                        # Читаем DOCX файлы
                        elif file_ext in docx_extensions:
                            if not DOCX_AVAILABLE or Document is None:
                                logger.warning(f"[RAG] python-docx not installed, skipping DOCX: {file_path.name}")
                                continue
                            
                            doc = Document(file_path)
                            text_parts: List[str] = []
                            for paragraph in doc.paragraphs:
                                if paragraph.text.strip():
                                    text_parts.append(paragraph.text)
                            content = "\n".join(text_parts)
                        
                        else:
                            continue
                        
                        if content:
                            # Разбиваем на чанки
                            file_chunks = GeminiService._split_text_into_chunks(content, chunk_size=1000, overlap=200)
                            common_chunks.extend(file_chunks)
                            common_metadata.extend([{"filename": f"common/{file_path.name}"} for _ in file_chunks])
                            logger.info(f"[RAG] Processed common/{file_path.name}: {len(file_chunks)} chunks")
                    
                    except Exception as e:
                        logger.warning(f"[RAG] Failed to process common/{file_path.name}: {e}")
            
            logger.info(f"[RAG] Common knowledge: {len(common_chunks)} chunks")
            
            # Теперь создаем индекс для каждого отдела
            for department in departments:
                try:
                    logger.info(f"[RAG] Creating index for department: {department}")
                    
                    dept_chunks: List[str] = []
                    dept_metadata: List[dict] = []
                    
                    # Добавляем common чанки
                    dept_chunks.extend(common_chunks)
                    dept_metadata.extend(common_metadata)
                    
                    # Читаем файлы отдела
                    dept_path = knowledge_path / department
                    if dept_path.exists() and dept_path.is_dir():
                        for file_path in dept_path.rglob("*"):
                            if not file_path.is_file():
                                continue
                            
                            file_ext = file_path.suffix.lower()
                            content = ""
                            
                            try:
                                if file_ext in text_extensions:
                                    content = file_path.read_text(encoding="utf-8")
                                elif file_ext in pdf_extensions:
                                    if fitz is None:
                                        continue
                                    doc = fitz.open(file_path)
                                    content = "\n".join([page.get_text() for page in doc])
                                    doc.close()
                                elif file_ext in docx_extensions:
                                    if not DOCX_AVAILABLE or Document is None:
                                        continue
                                    doc = Document(file_path)
                                    content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                                else:
                                    continue
                                
                                if content:
                                    file_chunks = GeminiService._split_text_into_chunks(content, chunk_size=1000, overlap=200)
                                    dept_chunks.extend(file_chunks)
                                    dept_metadata.extend([{"filename": f"{department}/{file_path.name}"} for _ in file_chunks])
                                    logger.info(f"[RAG] Processed {department}/{file_path.name}: {len(file_chunks)} chunks")
                            
                            except Exception as e:
                                logger.warning(f"[RAG] Failed to process {department}/{file_path.name}: {e}")
                    
                    # Создаем индекс для отдела
                    if dept_chunks:
                        logger.info(f"[RAG] Department {department}: {len(dept_chunks)} chunks total")
                        embeddings = GeminiService._generate_embeddings(dept_chunks)
                        
                        if len(embeddings) != len(dept_chunks):
                            dept_chunks = dept_chunks[:embeddings.shape[0]]
                            dept_metadata = dept_metadata[:embeddings.shape[0]]
                        
                        vector_store = VectorStore()
                        vector_store.clear()
                        vector_store.add_embeddings(embeddings, dept_chunks, dept_metadata)
                        GeminiService._vector_stores[department] = vector_store
                        logger.info(f"[RAG] Index created for {department}: {len(dept_chunks)} chunks")
                    else:
                        logger.warning(f"[RAG] No chunks for department {department}")
                
                except Exception as e:
                    logger.error(f"[RAG] Error creating index for {department}: {e}", exc_info=True)
            
            logger.info(f"[RAG] Created {len(GeminiService._vector_stores)} department indices")
            
            # Fallback: создаем старый глобальный индекс для обратной совместимости
            if common_chunks:
                all_chunks = common_chunks.copy()
                all_metadata = common_metadata.copy()
                embeddings = GeminiService._generate_embeddings(all_chunks)
                if len(embeddings) != len(all_chunks):
                    all_chunks = all_chunks[:embeddings.shape[0]]
                    all_metadata = all_metadata[:embeddings.shape[0]]
                vector_store = VectorStore()
                vector_store.clear()
                vector_store.add_embeddings(embeddings, all_chunks, all_metadata)
                GeminiService._vector_store = vector_store
                logger.info(f"[RAG] Fallback global index created with {len(all_chunks)} chunks")
            
        except Exception as e:
            logger.error(f"[RAG] Error creating vector index: {e}", exc_info=True)
            GeminiService._vector_store = None
    
    @staticmethod
    def _load_knowledge_base() -> str:
        """
        Загружает содержимое файлов из папки data/knowledge.
        
        Returns:
            Объединенное содержимое всех текстовых файлов из папки knowledge
        """
        knowledge_path = Path("data/knowledge")
        if not knowledge_path.exists():
            logger.info("Knowledge base directory not found")
            return ""
        
        context_parts: List[str] = []
        
        # Поддерживаемые текстовые форматы
        text_extensions = {".txt", ".md", ".rst"}
        pdf_extensions = {".pdf"}
        
        # Читаем все файлы из папки knowledge
        for file_path in knowledge_path.iterdir():
            if not file_path.is_file():
                continue
                
            file_ext = file_path.suffix.lower()
            content = ""
            
            try:
                # Читаем текстовые файлы
                if file_ext in text_extensions:
                    content = file_path.read_text(encoding="utf-8")
                    
                # Читаем PDF файлы
                elif file_ext in pdf_extensions:
                    if fitz is None:
                        logger.warning(f"PyMuPDF not installed, skipping PDF file: {file_path.name}")
                        continue
                    
                    doc = fitz.open(file_path)
                    text_parts: List[str] = []
                    for page in doc:
                        text_parts.append(page.get_text())
                    content = "\n".join(text_parts)
                    doc.close()
                    
                else:
                    logger.debug(f"Unsupported file type: {file_path.name}")
                    continue
                
                if content:
                    context_parts.append(f"Файл: {file_path.name}\n{content}\n")
                    logger.info(f"Loaded knowledge file: {file_path.name}")
                    
            except Exception as e:
                logger.warning(f"Failed to read file {file_path.name}: {e}")
        
        if not context_parts:
            logger.info("No text files found in knowledge base")
            return ""
        
        return "\n---\n".join(context_parts)
    
    @staticmethod
    def _split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Разрезает текст на куски с перекрытием.
        
        Args:
            text: Исходный текст
            chunk_size: Размер чанка в символах (1000-1500)
            overlap: Перекрытие между чанками в символах (200)
        
        Returns:
            Список текстовых чанков
        """
        if not text:
            return []
        
        chunks: List[str] = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            chunk = text[start:end]
            
            # Если не последний чанк, пытаемся закончить на границе предложения
            if end < text_length:
                # Ищем последнюю точку, восклицательный или вопросительный знак
                last_sentence_end = max(
                    chunk.rfind('.'),
                    chunk.rfind('!'),
                    chunk.rfind('?'),
                    chunk.rfind('\n')
                )
                if last_sentence_end > chunk_size * 0.5:  # Если нашли в последней половине
                    chunk = chunk[:last_sentence_end + 1]
                    end = start + last_sentence_end + 1
            
            chunks.append(chunk.strip())
            
            # Следующий чанк начинается с перекрытием
            start = end - overlap
            if start >= text_length:
                break
        
        logger.info(f"Split text into {len(chunks)} chunks (size: {chunk_size}, overlap: {overlap})")
        return chunks
    
    @staticmethod
    def _generate_embeddings(texts: List[str]) -> np.ndarray:
        """
        Генерирует эмбеддинги для списка текстов через новый API google.genai.
        
        Args:
            texts: Список текстов для эмбеддинга
        
        Returns:
            Массив эмбеддингов (numpy array, shape: [n_texts, dimension])
        """
        try:
            if gemini_client is None:
                raise ValueError("Gemini client not initialized")
            
            embeddings_list: List[List[float]] = []
            
            # Генерируем эмбеддинги для каждого текста через новый API
            for i, text in enumerate(texts):
                try:
                    # Используем gemini-embedding-001 для генерации эмбеддингов
                    result = gemini_client.models.embed_content(
                        model="gemini-embedding-001",
                        contents=text,
                        config=types.EmbedContentConfig(
                            task_type="RETRIEVAL_DOCUMENT"
                        )
                    )
                    
                    # Извлекаем эмбеддинг из результата
                    if hasattr(result, 'embeddings') and len(result.embeddings) > 0:
                        embedding = result.embeddings[0].values
                        embeddings_list.append(embedding)
                    else:
                        logger.warning(f"No embeddings in result for text {i}")
                        continue
                    
                    if (i + 1) % 10 == 0:
                        logger.info(f"Generated embeddings for {i + 1}/{len(texts)} texts")
                        
                except Exception as e:
                    logger.error(f"Error generating embedding for text {i}: {e}")
                    # Пропускаем этот текст вместо добавления нулевого вектора
                    logger.warning(f"Skipping text {i} due to embedding error")
                    continue
            
            if not embeddings_list:
                raise ValueError("Не удалось сгенерировать ни одного эмбеддинга")
            
            # Фильтруем чанки, для которых не удалось сгенерировать эмбеддинги
            if len(embeddings_list) != len(texts):
                logger.warning(f"Generated {len(embeddings_list)} embeddings for {len(texts)} texts (some were skipped)")
            
            embeddings_array = np.array(embeddings_list, dtype=np.float32)
            logger.info(f"Generated {len(embeddings_list)} embeddings (shape: {embeddings_array.shape})")
            return embeddings_array
            
        except Exception as e:
            logger.error(f"Error in _generate_embeddings: {e}", exc_info=True)
            raise Exception(f"Ошибка при генерации эмбеддингов: {str(e)}")
    
    @staticmethod
    def create_vector_db() -> None:
        """
        Создает векторную базу данных из файлов knowledge.
        Читает все файлы, разбивает на чанки, генерирует эмбеддинги и сохраняет в FAISS.
        """
        try:
            logger.info("[VECTOR_DB] Starting vector database creation...")
            
            knowledge_path = Path("data/knowledge")
            if not knowledge_path.exists():
                logger.warning("[VECTOR_DB] Knowledge base directory not found")
                return
            
            # Поддерживаемые форматы
            text_extensions = {".txt", ".md", ".rst"}
            pdf_extensions = {".pdf"}
            docx_extensions = {".docx"}
            
            all_chunks: List[str] = []
            
            # Читаем все файлы из папки knowledge
            for file_path in knowledge_path.iterdir():
                if not file_path.is_file():
                    continue
                
                file_ext = file_path.suffix.lower()
                content = ""
                
                try:
                    # Читаем текстовые файлы
                    if file_ext in text_extensions:
                        content = file_path.read_text(encoding="utf-8")
                    
                    # Читаем PDF файлы
                    elif file_ext in pdf_extensions:
                        if fitz is None:
                            logger.warning(f"[VECTOR_DB] PyMuPDF not installed, skipping PDF: {file_path.name}")
                            continue
                        
                        doc = fitz.open(file_path)
                        text_parts: List[str] = []
                        for page in doc:
                            text_parts.append(page.get_text())
                        content = "\n".join(text_parts)
                        doc.close()
                    
                    # Читаем DOCX файлы
                    elif file_ext in docx_extensions:
                        if not DOCX_AVAILABLE or Document is None:
                            logger.warning(f"[VECTOR_DB] python-docx not installed, skipping DOCX: {file_path.name}")
                            continue
                        
                        doc = Document(file_path)
                        text_parts: List[str] = []
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                text_parts.append(paragraph.text)
                        content = "\n".join(text_parts)
                    
                    else:
                        logger.debug(f"[VECTOR_DB] Unsupported file type: {file_path.name}")
                        continue
                    
                    if content:
                        # Разбиваем на чанки
                        file_chunks = GeminiService._split_text_into_chunks(
                            content,
                            chunk_size=1200,
                            overlap=200
                        )
                        
                        # Добавляем метаинформацию о файле к каждому чанку
                        file_chunks_with_meta = [
                            f"[Файл: {file_path.name}]\n{chunk}"
                            for chunk in file_chunks
                        ]
                        
                        all_chunks.extend(file_chunks_with_meta)
                        logger.info(f"[VECTOR_DB] Processed {file_path.name}: {len(file_chunks)} chunks")
                    
                except Exception as e:
                    logger.warning(f"[VECTOR_DB] Failed to process {file_path.name}: {e}")
            
            if not all_chunks:
                logger.warning("[VECTOR_DB] No chunks to process")
                return
            
            logger.info(f"[VECTOR_DB] Total chunks: {len(all_chunks)}")
            
            # Генерируем эмбеддинги
            logger.info("[VECTOR_DB] Generating embeddings...")
            embeddings = GeminiService._generate_embeddings(all_chunks)
            
            # Фильтруем чанки, для которых не удалось сгенерировать эмбеддинги
            # (если были пропущены некоторые тексты)
            if embeddings.shape[0] < len(all_chunks):
                logger.warning(f"[VECTOR_DB] Some chunks were skipped: {embeddings.shape[0]} embeddings for {len(all_chunks)} chunks")
                # Используем только те чанки, для которых есть эмбеддинги
                all_chunks = all_chunks[:embeddings.shape[0]]
            
            # Создаем и сохраняем векторное хранилище
            vector_store = VectorStore()
            vector_store.clear()  # Очищаем старые данные
            vector_store.add_embeddings(embeddings, all_chunks)
            vector_store.save_index()
            
            logger.info(f"[VECTOR_DB] Vector database created successfully with {len(all_chunks)} chunks")
            
        except Exception as e:
            logger.error(f"[VECTOR_DB] Error creating vector database: {e}", exc_info=True)
            raise Exception(f"Ошибка при создании векторной базы: {str(e)}")
    
    @staticmethod
    def _is_russian_text(text: str) -> bool:
        """
        Проверяет, содержит ли текст преимущественно русские символы (кириллицу).
        
        Args:
            text: Текст для проверки
            
        Returns:
            True если текст на русском, False иначе
        """
        if not text:
            return True
        
        # Простая проверка на наличие кириллицы
        return any('а' <= char <= 'я' or 'А' <= char <= 'Я' for char in text)
    
    @staticmethod
    async def _translate_to_russian(text: str) -> str:
        """
        Переводит текст на русский язык через Gemini для улучшения поиска в RAG.
        
        Args:
            text: Текст для перевода
            
        Returns:
            Переведенный текст на русском языке
        """
        try:
            if gemini_client is None:
                logger.warning("[TRANSLATE] Gemini client not initialized, returning original text")
                return text
            
            logger.info(f"[TRANSLATE] Translating query to Russian: {text[:100]}...")
            
            response = gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=f"Переведи этот текст на русский язык одной фразой, сохраняя смысл:\n\n{text}",
                config=types.GenerateContentConfig(
                    temperature=0.3,
                )
            )
            
            translated = response.text.strip() if response.text else text
            logger.info(f"[TRANSLATE] Translated query: {translated}")
            return translated
            
        except Exception as e:
            logger.error(f"[TRANSLATE] Error translating text: {e}", exc_info=True)
            return text  # Возвращаем оригинал при ошибке
    
    @staticmethod
    async def get_answer(
        prompt: str,
        user_id: int,
        session: AsyncSession,
        context: str | None = None,
    ) -> str:
        """
        Генерирует ответ на основе промпта и контекста из RAG-системы с использованием векторного поиска.
        Использует историю диалога для контекста.
        
        Args:
            prompt: Вопрос пользователя
            user_id: Telegram ID пользователя
            session: Сессия базы данных для работы с историей
            context: Дополнительный контекст (если передан, используется вместо векторного поиска)
        
        Returns:
            Сгенерированный ответ от модели Gemini
        
        Raises:
            Exception: Если не удалось сгенерировать ответ
        """
        try:
            # Проверяем наличие API ключа
            if not settings.gemini_api_key:
                raise ValueError("GEMINI_API_KEY не установлен. Установите его в .env файле.")
            
            # Получаем историю диалога (последние 10 сообщений)
            history_messages = await get_recent_messages(session, user_id, limit=10)
            history_text = format_history_for_prompt(history_messages)
            
            # Если контекст передан явно, используем его (для обратной совместимости)
            source_files: List[str] = []
            if context is not None:
                relevant_chunks_text = context
                source_files = []  # При явном контексте источники не определяем
            else:
                # Используем векторный поиск с учетом отдела пользователя
                try:
                    # Получаем отдел пользователя для изоляции знаний
                    from app.utils.department import get_user_department
                    user_department = await get_user_department(session, user_id)
                    
                    logger.info(f"[RAG] User {user_id} department: {user_department or 'admin (all departments)'}")
                    
                    # Проверяем, создан ли индекс
                    if not GeminiService._vector_stores:
                        logger.info("[RAG] Vector indices not found, creating new ones...")
                        GeminiService._create_department_indices()
                    
                    # Проверяем язык запроса и переводим на русский для точного поиска
                    search_query = prompt
                    if not GeminiService._is_russian_text(prompt):
                        logger.info(f"[RAG] Query is not in Russian, translating for better search accuracy...")
                        search_query = await GeminiService._translate_to_russian(prompt)
                    
                    # Генерируем эмбеддинг для запроса через новый API
                    logger.info(f"[RAG] Generating query embedding for: {search_query[:100]}...")
                    if gemini_client is None:
                        raise ValueError("Gemini client not initialized")
                    
                    query_embedding_result = gemini_client.models.embed_content(
                        model="gemini-embedding-001",
                        contents=search_query,
                        config=types.EmbedContentConfig(
                            task_type="RETRIEVAL_QUERY"
                        )
                    )
                    query_embedding = np.array(query_embedding_result.embeddings[0].values, dtype=np.float32)
                    
                    # РЕЖИМ БОГА ДЛЯ АДМИНА: Ищем по ВСЕМ индексам
                    if user_department is None:
                        logger.info(f"[RAG] 🔥 ADMIN GOD MODE: Searching across ALL department indices...")
                        all_search_results = []
                        
                        # Ищем по всем индексам отделов
                        for dept_name, dept_store in GeminiService._vector_stores.items():
                            if dept_store and dept_store.index is not None:
                                try:
                                    dept_results = dept_store.search(query_embedding, top_k=2)  # По 2 из каждого
                                    for chunk, distance, metadata in dept_results:
                                        # Добавляем информацию об отделе в метаданные
                                        enhanced_metadata = metadata.copy() if metadata else {}
                                        enhanced_metadata['department'] = dept_name
                                        all_search_results.append((chunk, distance, enhanced_metadata))
                                except Exception as e:
                                    logger.warning(f"[RAG] Error searching in {dept_name}: {e}")
                        
                        # Сортируем по distance и берем top-5 лучших
                        all_search_results.sort(key=lambda x: x[1])  # Меньше distance = лучше
                        search_results = all_search_results[:5]
                        logger.info(f"[RAG] Admin found {len(search_results)} chunks across {len(GeminiService._vector_stores)} departments")
                    
                    # Обычный режим для пользователей отделов
                    elif user_department and user_department in GeminiService._vector_stores:
                        # Пользователь с конкретным отделом ищет в своем отделе + common
                        search_results = []
                        
                        # 1. Поиск в своем отделе (приоритет)
                        vector_store = GeminiService._vector_stores[user_department]
                        logger.info(f"[RAG] User {user_id} (Dept: {user_department}) searching in department index...")
                        dept_results = vector_store.search(query_embedding, top_k=2)
                        search_results.extend(dept_results)
                        
                        # 2. Поиск в common (если он существует и это не сам common)
                        if user_department != "common" and "common" in GeminiService._vector_stores:
                            common_store = GeminiService._vector_stores["common"]
                            logger.info(f"[RAG] Also searching in 'common' for user {user_id}...")
                            common_results = common_store.search(query_embedding, top_k=2)
                            # Добавляем метаданные чтобы знать что из common
                            for chunk, distance, metadata in common_results:
                                enhanced_metadata = metadata.copy() if metadata else {}
                                enhanced_metadata['department'] = 'common'
                                search_results.append((chunk, distance, enhanced_metadata))
                        
                        # Сортируем по relevance (distance)
                        search_results.sort(key=lambda x: x[1])
                        search_results = search_results[:3]  # Топ-3 из обоих источников
                        
                        logger.info(f"[RAG] Found {len(search_results)} chunks (from {user_department} + common)")
                    
                    else:
                        logger.warning(f"[RAG] Department {user_department} not found in indices, using fallback")
                        vector_store = GeminiService._vector_store
                        if vector_store and vector_store.index is not None:
                            search_results = vector_store.search(query_embedding, top_k=3)
                        else:
                            search_results = []
                    
                    # Обработка результатов поиска
                    if not search_results:
                        logger.warning("[RAG] No relevant chunks found, using empty context")
                        relevant_chunks_text = ""
                        source_files = []
                    else:
                        # ДЕДУПЛИКАЦИЯ для админа (чтобы common/ не дублировался)
                        chunks_texts: List[str] = []
                        source_files_set: set[str] = set()
                        seen_chunks: set[str] = set()  # Для отслеживания уникальных чанков
                        departments_used: set[str] = set()  # Для отслеживания использованных отделов
                        
                        for chunk, distance, metadata in search_results:
                            # Создаем хеш чанка для проверки уникальности
                            chunk_hash = chunk.strip()[:200]  # Первые 200 символов для идентификации
                            
                            # Если админ - проверяем на дубликаты
                            if user_department is None:
                                if chunk_hash in seen_chunks:
                                    logger.debug(f"[RAG] Skipping duplicate chunk from {metadata.get('filename', 'unknown')}")
                                    continue  # Пропускаем дубликат
                                seen_chunks.add(chunk_hash)
                                
                                # МЕТКИ ИСТОЧНИКОВ ДЛЯ АДМИНА
                                if metadata and "filename" in metadata:
                                    filename = metadata.get("filename", "")
                                    department_name = metadata.get("department", "unknown")
                                    
                                    # ПРИОРИТЕТ ИСТОЧНИКОВ:
                                    # Если файл из common/ -> "Общие знания"
                                    if "common/" in filename or filename.startswith("common/"):
                                        source_label = "Общие знания"
                                    else:
                                        # Иначе показываем отдел
                                        source_label = department_name
                                    
                                    departments_used.add(source_label)
                                    # Добавляем метку источника перед чанком
                                    tagged_chunk = f"[Источник: {source_label}]\n{chunk}"
                                    chunks_texts.append(tagged_chunk)
                                else:
                                    chunks_texts.append(chunk)
                            else:
                                # Обычный пользователь - без меток
                                chunks_texts.append(chunk)
                            
                            # Извлекаем имя файла из метаданных
                            if metadata and "filename" in metadata:
                                filename = metadata["filename"]
                                # Для админа показываем отдел
                                if user_department is None and "department" in metadata:
                                    filename = f"[{metadata['department']}] {filename}"
                                source_files_set.add(filename)
                        
                        relevant_chunks_text = "\n\n---\n\n".join(chunks_texts)
                        source_files = sorted(list(source_files_set))
                        
                        if user_department is None:
                            logger.info(f"[RAG] 🔥 Admin: Found {len(chunks_texts)} unique chunks (after deduplication) from {len(source_files)} files")
                            logger.info(f"[RAG] 🏢 Departments used: {sorted(departments_used)}")
                        else:
                            logger.info(f"[RAG] Found {len(search_results)} relevant chunks from {len(source_files)} files: {source_files}")
                
                except Exception as e:
                    logger.error(f"[RAG] Error in vector search: {e}", exc_info=True)
                    # Fallback на старый метод
                    logger.info("[RAG] Falling back to full text search...")
                    relevant_chunks_text = GeminiService._load_knowledge_base()
                    source_files = []  # В fallback режиме источники не определяем
            
            # Системная инструкция с поддержкой многоязычности
            # Определяем, админ ли пользователь
            is_admin = user_department is None if 'user_department' in locals() else False
            
            system_instruction = """Ты — помощник UQsoft. Твоя задача — помогать сотрудникам находить информацию в базе знаний компании.

КРИТИЧНО - ЯЗЫК ОТВЕТА:
Определяй язык вопроса пользователя и отвечай строго на том же языке (русский, английский или китайский). 
Используй информацию из предоставленного контекста, но переводи её на язык запроса, если это необходимо.

Примеры:
- Вопрос на русском → Ответ на русском (переводи контекст если он на другом языке)
- Question in English → Answer in English (translate context if needed)
- 中文问题 → 中文回答 (翻译上下文如果需要)

ВАЖНО: Если не уверен в языке вопроса — используй русский язык по умолчанию."""
            
            # ОБЯЗАТЕЛЬНЫЙ ФОРМАТ ДЛЯ АДМИНА
            if is_admin:
                system_instruction += """

📂 ФОРМАТ ДЛЯ АДМИНИСТРАТОРА (ОБЯЗАТЕЛЬНО):
Ты видишь теги [Источник: ...] в контексте. Это означает что информация взята из разных отделов компании.
В КОНЦЕ своего ответа (с новой строки) ОБЯЗАТЕЛЬНО напиши:

📂 **Источники:** [список отделов через запятую]

Пример:
[Источник: Общие знания] Пароль Wi-Fi...
[Источник: sorting] Код сортировки...

Твой ответ ДОЛЖЕН заканчиваться на:

📂 **Источники:** Общие знания, sorting"""
            else:
                system_instruction += """

МЕТКИ ИСТОЧНИКОВ:
Если ты видишь теги [Источник: ...], это означает информация из разных источников.
Можешь упомянуть источники в конце ответа: "📚 Источники: источник1, источник2" """
            
            # Формируем промпт с контекстом, историей и вопросом
            prompt_parts: List[str] = []
            
            # Добавляем историю диалога, если она есть
            if history_text:
                prompt_parts.append(f"История предыдущего диалога:\n{history_text}\n")
            
            # Добавляем контекст из базы знаний
            if relevant_chunks_text and source_files:
                # Если есть источники, добавляем инструкцию об их отображении
                sources_text = ", ".join(source_files)
                prompt_parts.append(
                    "Используй предоставленные фрагменты знаний, чтобы ответить на вопрос. "
                    "Если ответа нет в тексте, так и скажи.\n\n"
                    f"Контекст:\n{relevant_chunks_text}\n\n"
                )
                prompt_parts.append(f"Вопрос: {prompt}\n\n")
                prompt_parts.append(
                    f"ВАЖНО: В конце ответа обязательно добавь список источников в формате:\n"
                    f"Источники: {sources_text}"
                )
            elif relevant_chunks_text:
                # Если есть контекст, но нет источников (старый формат)
                prompt_parts.append(
                    "Используй предоставленные фрагменты знаний, чтобы ответить на вопрос. "
                    "Если ответа нет в тексте, так и скажи.\n\n"
                    f"Контекст:\n{relevant_chunks_text}\n\n"
                    f"Вопрос: {prompt}"
                )
            else:
                # Если контекста нет, источники не нужны
                prompt_parts.append(
                    "В базе знаний пока нет информации.\n\n"
                    f"Вопрос: {prompt}"
                )
            
            full_prompt = "\n".join(prompt_parts)
            
            logger.info(f"[GEMINI] Generating response with Gemini for prompt: {prompt[:100]}...")
            
            # Проверяем наличие клиента
            if gemini_client is None:
                raise ValueError("Gemini client not initialized")
            
            # Генерируем ответ через новый API
            logger.info("[GEMINI] Generating content with gemini-2.5-flash...")
            response = gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=full_prompt,
                config=types.GenerateContentConfig(
                    system_instruction=system_instruction,
                    temperature=0.7,
                )
            )
            
            # Получаем текст ответа
            response_text = response.text if response.text else "Извините, не удалось сгенерировать ответ."
            
            logger.info(f"[GEMINI] Successfully generated response (length: {len(response_text)})")
            
            # Сохраняем вопрос пользователя и ответ бота в историю
            try:
                await save_message(session, user_id, "user", prompt)
                await save_message(session, user_id, "assistant", response_text)
                logger.info(f"[CHAT_HISTORY] Saved question and answer for user_id={user_id}")
            except Exception as e:
                logger.error(f"[CHAT_HISTORY] Failed to save history: {e}", exc_info=True)
                # Не прерываем выполнение, если не удалось сохранить историю
            
            return response_text
        
        except ResourceExhausted as e:
            logger.warning(f"[GEMINI] Quota exceeded (429): {e}")
            return QUOTA_EXCEEDED_MESSAGE
            
        except Exception as e:
            logger.error(f"Error generating response with Gemini: {e}", exc_info=True)
            raise Exception(f"Ошибка при генерации ответа: {str(e)}")
    
    @staticmethod
    async def get_answer_from_audio_with_rag(
        audio_bytes: bytes | None = None,
        audio_file_path: str | None = None,
        audio_mime_type: str = "audio/ogg",
        user_id: int | None = None,
        session: AsyncSession | None = None
    ) -> str:
        """
        Генерирует ответ на основе аудио файла С RAG ПОИСКОМ.
        Двухэтапная обработка:
        1. Транскрипция аудио -> текст вопроса
        2. RAG поиск по тексту -> релевантный контекст
        3. Генерация ответа с контекстом
        
        Args:
            audio_bytes: Байты аудио файла
            audio_file_path: Путь к аудио файлу
            audio_mime_type: MIME тип (по умолчанию audio/ogg)
            user_id: ID пользователя для RAG поиска по отделу
            session: AsyncSession для работы с БД
        
        Returns:
            Сгенерированный ответ от модели Gemini
        """
        try:
            logger.info(f"[VOICE_RAG] Starting audio processing with RAG: file_path={audio_file_path}")
            
            if gemini_client is None:
                raise ValueError("Gemini client not initialized")
            
            # ЭТАП 1: ТРАНСКРИПЦИЯ АУДИО
            logger.info("[VOICE_RAG] Step 1: Transcribing audio to text...")
            transcribed_text = await GeminiService._transcribe_audio(
                audio_bytes=audio_bytes,
                audio_file_path=audio_file_path,
                audio_mime_type=audio_mime_type
            )
            logger.info(f"[VOICE_RAG] Transcribed text: {transcribed_text[:200]}...")
            
            # ЭТАП 2: RAG ПОИСК ПО ТЕКСТУ
            logger.info("[VOICE_RAG] Step 2: Performing RAG search...")
            if user_id and session:
                # Используем тот же RAG механизм что и в get_answer
                from app.utils.department import get_user_department
                user_department = await get_user_department(session, user_id)
                
                # Проверяем язык и переводим для точного поиска
                search_query = transcribed_text
                if not GeminiService._is_russian_text(transcribed_text):
                    logger.info("[VOICE_RAG] Translating query to Russian for better RAG accuracy...")
                    search_query = await GeminiService._translate_to_russian(transcribed_text)
                
                # Генерируем эмбеддинг
                query_embedding_result = gemini_client.models.embed_content(
                    model="gemini-embedding-001",
                    contents=search_query,
                    config=types.EmbedContentConfig(task_type="RETRIEVAL_QUERY")
                )
                query_embedding = np.array(query_embedding_result.embeddings[0].values, dtype=np.float32)
                
                # Поиск по индексам (God Mode для админа или по отделу)
                search_results = []
                departments_used: set[str] = set()
                
                if user_department is None:  # Админ - поиск везде
                    logger.info("[VOICE_RAG] 🔥 Admin God Mode for voice!")
                    for dept_name, dept_store in GeminiService._vector_stores.items():
                        if dept_store and dept_store.index:
                            try:
                                dept_results = dept_store.search(query_embedding, top_k=2)
                                for chunk, distance, metadata in dept_results:
                                    enhanced_metadata = metadata.copy() if metadata else {}
                                    enhanced_metadata['department'] = dept_name
                                    search_results.append((chunk, distance, enhanced_metadata))
                            except Exception as e:
                                logger.warning(f"[VOICE_RAG] Error searching {dept_name}: {e}")
                    search_results.sort(key=lambda x: x[1])
                    search_results = search_results[:5]
                elif user_department in GeminiService._vector_stores:
                    vector_store = GeminiService._vector_stores[user_department]
                    search_results = vector_store.search(query_embedding, top_k=3)
                
                # Формируем контекст с метками
                chunks_texts = []
                seen_chunks = set()
                for chunk, distance, metadata in search_results:
                    chunk_hash = chunk.strip()[:200]
                    if user_department is None:  # Дедупликация для админа
                        if chunk_hash in seen_chunks:
                            continue
                        seen_chunks.add(chunk_hash)
                        # Метки источников для админа
                        if metadata and "department" in metadata:
                            dept_name = metadata["department"]
                            departments_used.add(dept_name)
                            tagged_chunk = f"[Источник: {dept_name}]\n{chunk}"
                            chunks_texts.append(tagged_chunk)
                        else:
                            chunks_texts.append(chunk)
                    else:
                        chunks_texts.append(chunk)
                
                context = "\n\n---\n\n".join(chunks_texts)
                logger.info(f"[VOICE_RAG] RAG context prepared: {len(context)} chars, {len(chunks_texts)} chunks")
                if user_department is None:
                    logger.info(f"[VOICE_RAG] 🏢 Departments used: {sorted(departments_used)}")
            else:
                # Fallback если нет user_id/session
                logger.warning("[VOICE_RAG] No user_id/session, using fallback context")
                context = GeminiService._load_knowledge_base()
            
            # ЭТАП 3: ГЕНЕРАЦИЯ ОТВЕТА С RAG КОНТЕКСТОМ
            logger.info(f"[VOICE_RAG] Step 3: Generating answer with RAG context ({len(context)} chars)...")
            
            # Системная инструкция для голосовых сообщений
            is_admin = user_department is None if user_id and session else False
            
            system_instruction = """Ты — помощник UQsoft. Твоя задача — помогать сотрудникам находить информацию в базе знаний компании.

КРИТИЧНО - ЯЗЫК ОТВЕТА:
Определяй язык голосового сообщения пользователя и отвечай строго на том же языке (русский, английский или китайский). 
Используй информацию из предоставленного контекста, но переводи её на язык запроса, если это необходимо.

Пользователь спросил (голосом): {question}
Найди ответ в базе знаний."""
            
            # ОБЯЗАТЕЛЬНЫЙ ФОРМАТ ДЛЯ АДМИНА
            if is_admin:
                system_instruction += """

📂 ФОРМАТ ДЛЯ АДМИНИСТРАТОРА (ОБЯЗАТЕЛЬНО):
Ты видишь теги [Источник: ...] в контексте.
В КОНЦЕ своего ответа ОБЯЗАТЕЛЬНО напиши:

📂 **Источники:** [список отделов через запятую]"""
            
            # Формируем промпт
            prompt_parts = []
            if context:
                prompt_parts.append(f"=== КОНТЕКСТ ИЗ БАЗЫ ЗНАНИЙ ===\n{context}\n\n")
            prompt_parts.append(f"=== ВОПРОС ПОЛЬЗОВАТЕЛЯ ===\n{transcribed_text}")
            
            prompt = "\n".join(prompt_parts)
            
            # Генерируем ответ
            response = gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt,
                config=types.GenerateContentConfig(
                    system_instruction=system_instruction.replace("{question}", transcribed_text),
                    temperature=0.7,
                )
            )
            
            response_text = response.text if response.text else "Извините, не удалось сгенерировать ответ."
            logger.info(f"[VOICE_RAG] Successfully generated response (length: {len(response_text)})")
            
            return response_text
            
        except Exception as e:
            logger.error(f"[VOICE_RAG] Error: {e}", exc_info=True)
            raise Exception(f"Ошибка при обработке голосового сообщения: {str(e)}")
    
    @staticmethod
    async def _transcribe_audio(
        audio_bytes: bytes | None = None,
        audio_file_path: str | None = None,
        audio_mime_type: str = "audio/ogg"
    ) -> str:
        """
        Транскрибирует аудио в текст через Gemini.
        
        Returns:
            Распознанный текст
        """
        try:
            if gemini_client is None:
                raise ValueError("Gemini client not initialized")
            
            content_parts: List = []
            uploaded_file = None
            
            # Загружаем аудио файл
            if audio_file_path:
                try:
                    uploaded_file = gemini_client.files.upload(file=audio_file_path)
                    content_parts.append(uploaded_file)
                except Exception as upload_error:
                    logger.warning(f"[STT] Failed to upload: {upload_error}, using bytes")
                    with open(audio_file_path, "rb") as f:
                        audio_bytes = f.read()
            
            if not uploaded_file:
                if not audio_bytes:
                    raise ValueError("No audio data provided")
                content_parts.append({
                    "mime_type": audio_mime_type,
                    "data": audio_bytes
                })
            
            # Добавляем промпт для транскрипции
            content_parts.append("Распознай речь из этого аудио и верни ТОЛЬКО текст того, что сказал пользователь. Ничего кроме текста речи не пиши.")
            
            # Генерируем транскрипцию
            response = gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=content_parts,
                config=types.GenerateContentConfig(
                    temperature=0.3,
                )
            )
            
            # Очистка загруженного файла
            if uploaded_file:
                try:
                    gemini_client.files.delete(name=uploaded_file.name)
                except Exception:
                    pass
            
            transcribed_text = response.text if response.text else ""
            return transcribed_text.strip()
            
        except Exception as e:
            logger.error(f"[STT] Error transcribing audio: {e}", exc_info=True)
            raise
    
    @staticmethod
    def get_answer_from_audio(
        audio_bytes: bytes | None = None,
        audio_file_path: str | None = None,
        audio_mime_type: str = "audio/ogg",
        context: str | None = None
    ) -> str:
        """
        DEPRECATED: Старый метод без RAG. Используйте get_answer_from_audio_with_rag.
        
        Генерирует ответ на основе аудио файла (голосового сообщения) через новый API google.genai.
        
        Args:
            audio_bytes: Байты аудио файла (опционально, если передан audio_file_path)
            audio_file_path: Путь к аудио файлу (приоритет над audio_bytes)
            audio_mime_type: MIME тип аудио (по умолчанию audio/ogg для Telegram)
            context: Дополнительный контекст (если передан)
        
        Returns:
            Сгенерированный ответ от модели Gemini
        
        Raises:
            Exception: Если не удалось сгенерировать ответ
        """
        try:
            logger.info(f"[GEMINI] Starting audio processing: file_path={audio_file_path}, mime_type={audio_mime_type}")
            
            # Проверяем наличие клиента
            if gemini_client is None:
                raise ValueError("Gemini client not initialized")
            
            # Загружаем контекст из базы знаний, если не передан
            if context is None:
                logger.info("[GEMINI] Loading knowledge base context...")
                context = GeminiService._load_knowledge_base()
                logger.info(f"[GEMINI] Knowledge base loaded: {len(context)} chars")
            
            # Системная инструкция с многоязычностью
            system_instruction = """Ты — помощник UQsoft. Твоя задача — помогать сотрудникам находить информацию в базе знаний компании.

КРИТИЧНО - ЯЗЫК ОТВЕТА:
Определяй язык голосового сообщения пользователя и отвечай строго на том же языке (русский, английский или китайский). 
Используй информацию из предоставленного контекста, но переводи её на язык запроса, если это необходимо.

Пользователь прислал голосовое сообщение. Распознай речь, найди ответ в базе знаний и ответь на вопрос."""
            
            # Подготавливаем контент для отправки
            content_parts: List = []
            
            # Загружаем аудио файл через новый API
            uploaded_file = None
            if audio_file_path:
                try:
                    logger.info(f"[GEMINI] Uploading file via gemini_client.files.upload: {audio_file_path}")
                    uploaded_file = gemini_client.files.upload(file=audio_file_path)
                    logger.info(f"[GEMINI] File uploaded successfully: {uploaded_file.name}")
                    content_parts.append(uploaded_file)
                except Exception as upload_error:
                    logger.warning(f"[GEMINI] Failed to upload file: {upload_error}")
                    logger.info("[GEMINI] Falling back to direct bytes method...")
                    # Fallback: читаем файл в байты
                    with open(audio_file_path, "rb") as f:
                        audio_bytes = f.read()
                    logger.info(f"[GEMINI] Read file into bytes: {len(audio_bytes)} bytes")
            
            # Способ 2: Если не удалось загрузить через upload_file, используем байты
            if not uploaded_file:
                if not audio_bytes:
                    raise ValueError("Необходимо указать либо audio_file_path, либо audio_bytes")
                
                logger.info(f"[GEMINI] Using direct bytes method (size: {len(audio_bytes)} bytes)")
                
                # Пробуем отправить .ogg напрямую
                try:
                    content_parts.append({
                        "mime_type": audio_mime_type,
                        "data": audio_bytes
                    })
                    logger.info(f"[GEMINI] Added audio with mime_type={audio_mime_type}")
                except Exception as ogg_error:
                    logger.warning(f"[GEMINI] Failed to send .ogg directly: {ogg_error}")
                    
                    # Способ 3: Конвертируем .ogg в .wav через pydub
                    if PYDUB_AVAILABLE and audio_mime_type == "audio/ogg":
                        logger.info("[GEMINI] Attempting to convert .ogg to .wav using pydub...")
                        try:
                            import tempfile
                            temp_wav = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
                            temp_wav_path = temp_wav.name
                            temp_wav.close()
                            
                            # Конвертируем
                            if audio_file_path:
                                audio_segment = AudioSegment.from_file(audio_file_path, format="ogg")
                            else:
                                # Сохраняем байты во временный файл для конвертации
                                temp_ogg = tempfile.NamedTemporaryFile(suffix=".ogg", delete=False)
                                temp_ogg.write(audio_bytes)
                                temp_ogg_path = temp_ogg.name
                                temp_ogg.close()
                                
                                audio_segment = AudioSegment.from_file(temp_ogg_path, format="ogg")
                                Path(temp_ogg_path).unlink()  # Удаляем временный .ogg
                            
                            audio_segment.export(temp_wav_path, format="wav")
                            logger.info(f"[GEMINI] Converted to WAV: {temp_wav_path}")
                            
                            # Читаем .wav в байты
                            with open(temp_wav_path, "rb") as f:
                                wav_bytes = f.read()
                            
                            # Удаляем временный файл
                            Path(temp_wav_path).unlink()
                            
                            # Используем .wav
                            content_parts.append({
                                "mime_type": "audio/wav",
                                "data": wav_bytes
                            })
                            logger.info(f"[GEMINI] Using converted WAV (size: {len(wav_bytes)} bytes)")
                        except Exception as convert_error:
                            logger.error(f"[GEMINI] Failed to convert to WAV: {convert_error}", exc_info=True)
                            raise Exception(f"Не удалось обработать аудио файл: {str(convert_error)}")
                    else:
                        if not PYDUB_AVAILABLE:
                            logger.warning("[GEMINI] pydub not available, cannot convert audio")
                        raise Exception(f"Не удалось отправить аудио в Gemini: {str(ogg_error)}")
            
            # Добавляем текстовый промпт с контекстом
            prompt_text = "Распознай речь в аудио и ответь на вопрос пользователя, используя информацию из базы знаний."
            if context:
                prompt_text += f"\n\nКонтекст из базы знаний:\n{context}"
            
            content_parts.append(prompt_text)
            logger.info(f"[GEMINI] Added text prompt (length: {len(prompt_text)} chars)")
            
            logger.info(f"[GEMINI] Generating response with {len(content_parts)} content parts...")
            
            # Генерируем ответ через новый API
            response = gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=content_parts,
                config=types.GenerateContentConfig(
                    system_instruction=system_instruction,
                    temperature=0.7,
                )
            )
            
            # Получаем текст ответа
            response_text = response.text if response.text else "Извините, не удалось распознать речь или сгенерировать ответ."
            
            logger.info(f"[GEMINI] Successfully generated response (length: {len(response_text)} chars)")
            
            # Удаляем загруженный файл если использовали upload_file
            if uploaded_file:
                try:
                    gemini_client.files.delete(name=uploaded_file.name)
                    logger.info(f"[GEMINI] Cleaned up uploaded file: {uploaded_file.name}")
                except Exception as cleanup_error:
                    logger.warning(f"[GEMINI] Failed to cleanup uploaded file: {cleanup_error}")
            
            return response_text
        
        except ResourceExhausted as e:
            logger.warning(f"[GEMINI] Quota exceeded (429) for audio: {e}")
            return QUOTA_EXCEEDED_MESSAGE
            
        except Exception as e:
            logger.error(f"[GEMINI] ERROR generating response from audio: {type(e).__name__}: {e}", exc_info=True)
            raise Exception(f"Ошибка при обработке аудио: {str(e)}")
    
    @staticmethod
    def extract_media_links(text: str) -> dict[str, List[str]]:
        """
        Извлекает медиа-ссылки из текста (YouTube, файлы и т.д.).
        
        Args:
            text: Текст для анализа
        
        Returns:
            Словарь с типами ссылок и их значениями:
            {
                "youtube": ["https://youtube.com/watch?v=..."],
                "files": ["https://example.com/file.pdf"],
                "images": ["https://example.com/image.png"]
            }
        """
        import re
        
        media_links: dict[str, List[str]] = {
            "youtube": [],
            "files": [],
            "images": []
        }
        
        # Паттерны для поиска ссылок
        youtube_pattern = r'(?:https?://)?(?:www\.)?(?:youtube\.com/watch\?v=|youtu\.be/)([a-zA-Z0-9_-]{11})'
        file_pattern = r'https?://[^\s]+\.(?:pdf|doc|docx|xls|xlsx|ppt|pptx|zip|rar|txt|md)'
        image_pattern = r'https?://[^\s]+\.(?:jpg|jpeg|png|gif|webp|bmp|svg)'
        
        # Ищем YouTube ссылки
        youtube_matches = re.findall(youtube_pattern, text, re.IGNORECASE)
        for match in youtube_matches:
            full_url = f"https://www.youtube.com/watch?v={match}"
            if full_url not in media_links["youtube"]:
                media_links["youtube"].append(full_url)
        
        # Ищем файлы
        file_matches = re.findall(file_pattern, text, re.IGNORECASE)
        media_links["files"].extend(file_matches)
        
        # Ищем изображения
        image_matches = re.findall(image_pattern, text, re.IGNORECASE)
        media_links["images"].extend(image_matches)
        
        return media_links
    
    @staticmethod
    def process_knowledge_text(raw_text: str) -> tuple[str, str]:
        """
        Обрабатывает текст для добавления в базу знаний.
        Генерирует название файла и структурирует текст.
        
        Args:
            raw_text: Исходный текст от пользователя
        
        Returns:
            Кортеж (filename, structured_text):
            - filename: Короткое английское название файла (без расширения)
            - structured_text: Структурированный текст
        
        Raises:
            Exception: Если не удалось обработать текст
        """
        try:
            logger.info(f"[GEMINI] Processing knowledge text (length: {len(raw_text)} chars)...")
            
            # Проверяем наличие API ключа
            if not settings.gemini_api_key:
                raise ValueError("GEMINI_API_KEY не установлен. Установите его в .env файле.")
            
            # Проверяем наличие клиента
            if gemini_client is None:
                raise ValueError("Gemini client not initialized")
            
            # Системная инструкция для генерации названия и структурирования
            system_instruction = """Ты — AI-редактор базы знаний компании UQsoft.
Твоя задача:
1. Придумать короткое английское название файла (snake_case, без расширения, максимум 3 слова)
2. Структурировать текст: добавить заголовки, списки, выделить ключевые моменты
3. ВАЖНО: Сохрани оригинальный язык текста (не переводи контент)

ФОРМАТ ОТВЕТА (строго соблюдай):
FILENAME: название_файла
---
Структурированный текст здесь..."""
            
            # Формируем промпт
            prompt = f"Обработай следующий текст для базы знаний:\n\n{raw_text}"
            
            # Генерируем ответ через новый API
            logger.info("[GEMINI] Generating structured text with gemini-2.5-flash...")
            response = gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt,
                config=types.GenerateContentConfig(
                    system_instruction=system_instruction,
                    temperature=0.5,
                )
            )
            
            if not response.text:
                raise ValueError("Gemini не вернул ответ")
            
            response_text = response.text.strip()
            logger.info(f"[GEMINI] Received response (length: {len(response_text)} chars)")
            
            # Парсим ответ
            if "FILENAME:" not in response_text or "---" not in response_text:
                # Если формат не соблюден, генерируем название из первых слов
                import re
                words = re.findall(r'\b[a-zA-Zа-яА-Я]+\b', raw_text[:100])[:3]
                filename = "_".join(words).lower()[:30] if words else "knowledge_doc"
                structured_text = response_text
            else:
                # Парсим FILENAME и текст
                parts = response_text.split("---", 1)
                filename_line = parts[0].strip()
                structured_text = parts[1].strip() if len(parts) > 1 else raw_text
                
                # Извлекаем filename
                if "FILENAME:" in filename_line:
                    filename = filename_line.replace("FILENAME:", "").strip()
                else:
                    filename = "knowledge_doc"
            
            # Очищаем filename от недопустимых символов
            import re
            filename = re.sub(r'[^\w\-]', '_', filename).lower()
            filename = re.sub(r'_+', '_', filename).strip('_')[:50]
            
            if not filename:
                filename = "knowledge_doc"
            
            logger.info(f"[GEMINI] Generated filename: {filename}")
            
            return filename, structured_text
        
        except ResourceExhausted as e:
            logger.warning(f"[GEMINI] Quota exceeded (429) for knowledge processing: {e}")
            raise Exception(QUOTA_EXCEEDED_MESSAGE)
            
        except Exception as e:
            logger.error(f"[GEMINI] Error processing knowledge text: {e}", exc_info=True)
            raise Exception(f"Ошибка при обработке текста: {str(e)}")
    
    @staticmethod
    async def process_knowledge_audio(audio_path: Path) -> tuple[str, str]:
        """
        Обрабатывает голосовое сообщение для добавления в базу знаний.
        Использует Gemini 2.5 Flash Native Audio для транскрибации и структурирования.
        
        Args:
            audio_path: Путь к аудио-файлу (.ogg)
        
        Returns:
            Кортеж (filename, structured_text):
            - filename: Короткое английское название файла (без расширения)
            - structured_text: Структурированный текст из аудио
        
        Raises:
            Exception: Если не удалось обработать аудио
        """
        try:
            logger.info(f"[GEMINI] Processing knowledge audio: {audio_path.name}...")
            
            # Проверяем наличие клиента
            if gemini_client is None:
                raise ValueError("Gemini client not initialized")
            
            # Загружаем аудио-файл
            logger.info(f"[GEMINI] Uploading audio file to Gemini...")
            uploaded_file = gemini_client.files.upload(file=str(audio_path))
            logger.info(f"[GEMINI] Audio file uploaded: {uploaded_file.name}")
            
            # Системная инструкция для обработки аудио
            system_instruction = """Ты — аналитик UQsoft. Преврати это аудио в официальную базу знаний.
Удали мусор, структурируй по пунктам:
• Суть
• Детали
• Теги
• Задачи

ВАЖНО: Сохрани оригинальный язык аудио (не переводи контент).

ФОРМАТ ОТВЕТА (строго соблюдай):
FILENAME: короткое_название
---
# Суть
[основная мысль]

# Детали
[подробности]

# Теги
[ключевые слова]

# Задачи
[действия, если есть]
"""
            
            # Генерируем ответ через новый API с multimodal support (gemini-2.5-flash поддерживает аудио)
            logger.info("[GEMINI] Generating structured text from audio with gemini-2.5-flash...")
            response = gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[
                    "Проанализируй это аудио и создай структурированное знание для базы данных.",
                    uploaded_file,
                ],
                config=types.GenerateContentConfig(
                    system_instruction=system_instruction,
                    temperature=0.5,
                )
            )
            
            if not response.text:
                raise ValueError("Gemini не вернул ответ из аудио")
            
            response_text = response.text.strip()
            logger.info(f"[GEMINI] Received structured text from audio (length: {len(response_text)} chars)")
            
            # Парсим ответ
            if "FILENAME:" not in response_text or "---" not in response_text:
                # Если формат не соблюден, генерируем название
                import re
                from datetime import datetime
                timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                filename = f"audio_knowledge_{timestamp}"
                structured_text = response_text
            else:
                # Парсим FILENAME и текст
                parts = response_text.split("---", 1)
                filename_line = parts[0].strip()
                structured_text = parts[1].strip() if len(parts) > 1 else response_text
                
                # Извлекаем filename
                if "FILENAME:" in filename_line:
                    filename = filename_line.replace("FILENAME:", "").strip()
                else:
                    from datetime import datetime
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                    filename = f"audio_knowledge_{timestamp}"
            
            # Очищаем filename от недопустимых символов
            import re
            filename = re.sub(r'[^\w\-]', '_', filename).lower()
            filename = re.sub(r'_+', '_', filename).strip('_')[:50]
            
            if not filename:
                from datetime import datetime
                timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                filename = f"audio_knowledge_{timestamp}"
            
            logger.info(f"[GEMINI] Generated filename from audio: {filename}")
            
            # Удаляем загруженный файл из Gemini
            try:
                gemini_client.files.delete(name=uploaded_file.name)
                logger.info(f"[GEMINI] Deleted uploaded audio file: {uploaded_file.name}")
            except Exception as e:
                logger.warning(f"[GEMINI] Failed to delete uploaded file: {e}")
            
            return filename, structured_text
        
        except ResourceExhausted as e:
            logger.warning(f"[GEMINI] Quota exceeded (429) for audio processing: {e}")
            raise Exception(QUOTA_EXCEEDED_MESSAGE)
            
        except Exception as e:
            logger.error(f"[GEMINI] Error processing knowledge audio: {e}", exc_info=True)
            raise Exception(f"Ошибка при обработке аудио: {str(e)}")
    
    @staticmethod
    def get_knowledge_files() -> List[str]:
        """
        Возвращает список всех файлов из папки knowledge.
        
        Returns:
            Список имен файлов (без путей)
        """
        try:
            knowledge_path = Path("data/knowledge")
            if not knowledge_path.exists():
                logger.info("Knowledge base directory does not exist")
                return []
            
            files = [
                file.name
                for file in knowledge_path.iterdir()
                if file.is_file() and file.suffix in {".txt", ".pdf", ".md", ".rst"}
            ]
            
            logger.info(f"Found {len(files)} knowledge files")
            return sorted(files)
            
        except Exception as e:
            logger.error(f"Error getting knowledge files: {e}", exc_info=True)
            return []
    
    @staticmethod
    def delete_knowledge_file(filename: str) -> bool:
        """
        Удаляет файл из базы знаний.
        
        Args:
            filename: Имя файла для удаления (без пути)
        
        Returns:
            True если файл успешно удален, False если произошла ошибка
        
        Raises:
            Exception: Если файл не существует или произошла ошибка при удалении
        """
        try:
            knowledge_path = Path("data/knowledge")
            if not knowledge_path.exists():
                raise FileNotFoundError("Knowledge base directory does not exist")
            
            file_path = knowledge_path / filename
            
            # Проверяем, что файл существует и находится в правильной директории
            if not file_path.exists():
                raise FileNotFoundError(f"File {filename} does not exist in knowledge base")
            
            # Проверяем, что файл действительно в knowledge директории (защита от path traversal)
            if not str(file_path.resolve()).startswith(str(knowledge_path.resolve())):
                raise ValueError(f"Invalid file path: {filename}")
            
            # Удаляем файл
            file_path.unlink()
            
            logger.info(f"Deleted knowledge file: {filename}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting knowledge file {filename}: {e}", exc_info=True)
            raise Exception(f"Ошибка при удалении файла: {str(e)}")
    
    @staticmethod
    def get_knowledge_stats() -> dict[str, int]:
        """
        Собирает статистику по базе знаний - количество документов в каждом отделе.
        
        Returns:
            Словарь {department_name: file_count}, например:
            {
                "common": 5,
                "sorting": 12,
                "delivery": 8,
                "manager": 3,
                "customer_service": 7
            }
        """
        try:
            knowledge_path = Path("data/knowledge")
            if not knowledge_path.exists():
                logger.info("[STATS] Knowledge base directory does not exist")
                return {}
            
            # Поддерживаемые форматы файлов
            supported_extensions = {".txt", ".pdf", ".docx", ".md", ".rst"}
            
            stats: dict[str, int] = {}
            
            # Проходим по всем папкам (отделам) в data/knowledge
            for dept_path in knowledge_path.iterdir():
                if not dept_path.is_dir():
                    continue
                
                dept_name = dept_path.name
                
                # Считаем файлы рекурсивно (включая подпапки, например delivery/courier)
                file_count = 0
                for file_path in dept_path.rglob("*"):
                    if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                        file_count += 1
                
                if file_count > 0:
                    stats[dept_name] = file_count
                    logger.info(f"[STATS] Department '{dept_name}': {file_count} files")
            
            logger.info(f"[STATS] Total departments with files: {len(stats)}")
            return stats
            
        except Exception as e:
            logger.error(f"[STATS] Error getting knowledge stats: {e}", exc_info=True)
            return {}
    
    @staticmethod
    def get_department_files(dept_name: str) -> List[dict[str, str]]:
        """
        Получает список всех файлов в указанном отделе.
        
        Args:
            dept_name: Название отдела (например, 'sorting', 'manager')
        
        Returns:
            Список словарей с информацией о файлах:
            [
                {
                    "name": "guide.txt",
                    "path": "sorting/guide.txt",
                    "size": "1.2 KB",
                    "size_bytes": 1234
                },
                ...
            ]
        """
        try:
            knowledge_path = Path("data/knowledge")
            dept_path = knowledge_path / dept_name
            
            if not dept_path.exists() or not dept_path.is_dir():
                logger.warning(f"[FILES] Department '{dept_name}' not found")
                return []
            
            # Поддерживаемые форматы файлов
            supported_extensions = {".txt", ".pdf", ".docx", ".md", ".rst"}
            
            files_info: List[dict[str, str]] = []
            
            # Рекурсивно ищем все файлы в отделе
            for file_path in dept_path.rglob("*"):
                if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                    # Получаем относительный путь от knowledge/
                    relative_path = file_path.relative_to(knowledge_path)
                    
                    # Размер файла в байтах
                    size_bytes = file_path.stat().st_size
                    
                    # Форматируем размер файла
                    if size_bytes < 1024:
                        size_str = f"{size_bytes} B"
                    elif size_bytes < 1024 * 1024:
                        size_str = f"{size_bytes / 1024:.1f} KB"
                    else:
                        size_str = f"{size_bytes / (1024 * 1024):.1f} MB"
                    
                    files_info.append({
                        "name": file_path.name,
                        "path": str(relative_path).replace("\\", "/"),
                        "size": size_str,
                        "size_bytes": size_bytes
                    })
            
            # Сортируем по имени файла
            files_info.sort(key=lambda x: x["name"].lower())
            
            logger.info(f"[FILES] Found {len(files_info)} files in department '{dept_name}'")
            return files_info
            
        except Exception as e:
            logger.error(f"[FILES] Error getting files for department '{dept_name}': {e}", exc_info=True)
            return []
    
    @staticmethod
    def delete_document(dept_name: str, filename: str) -> bool:
        """
        Удаляет файл из базы знаний и обновляет векторные индексы.
        
        Args:
            dept_name: Название отдела (например, 'sorting', 'manager')
            filename: Имя файла для удаления
        
        Returns:
            True если файл успешно удален, False при ошибке
        
        Raises:
            FileNotFoundError: Если файл не существует
            ValueError: Если путь невалиден (защита от path traversal)
            Exception: При других ошибках
        """
        try:
            knowledge_path = Path("data/knowledge")
            if not knowledge_path.exists():
                raise FileNotFoundError("Knowledge base directory does not exist")
            
            # Безопасность: проверяем, что dept_name не содержит подозрительных символов
            if ".." in dept_name or "/" in dept_name or "\\" in dept_name:
                raise ValueError(f"Invalid department name: {dept_name}")
            
            dept_path = knowledge_path / dept_name
            if not dept_path.exists() or not dept_path.is_dir():
                raise FileNotFoundError(f"Department '{dept_name}' not found")
            
            # Ищем файл (может быть в подпапке, например delivery/courier/)
            file_path = None
            for found_file in dept_path.rglob(filename):
                if found_file.is_file() and found_file.name == filename:
                    file_path = found_file
                    break
            
            if not file_path:
                raise FileNotFoundError(f"File '{filename}' not found in department '{dept_name}'")
            
            # Проверяем, что файл действительно в knowledge директории (защита от path traversal)
            if not str(file_path.resolve()).startswith(str(knowledge_path.resolve())):
                raise ValueError(f"Invalid file path: {filename}")
            
            logger.info(f"[DELETE] Deleting file: {file_path}")
            
            # Удаляем файл
            file_path.unlink()
            logger.info(f"[DELETE] File deleted successfully: {file_path}")
            
            # Пересобираем индекс отдела
            try:
                logger.info(f"[DELETE] Rebuilding index for department: {dept_name}")
                GeminiService.rebuild_index_for_department(dept_name)
                logger.info(f"[DELETE] Index rebuilt successfully for department: {dept_name}")
            except Exception as rebuild_error:
                logger.error(f"[DELETE] Error rebuilding index: {rebuild_error}", exc_info=True)
                # Не прерываем процесс, файл уже удален
            
            return True
            
        except Exception as e:
            logger.error(f"[DELETE] Error deleting document '{filename}' from '{dept_name}': {e}", exc_info=True)
            raise


# Инициализация векторных индексов по отделам при загрузке модуля
try:
    GeminiService._create_department_indices()
except Exception as e:
    logger.error(f"[RAG] Failed to initialize department indices: {e}", exc_info=True)


# Оставляем функцию для обратной совместимости (устарела, требует user_id и session)
def generate_response(prompt: str, context: str | None = None) -> str:
    """
    Устаревшая функция. Используйте GeminiService.get_answer() вместо этого.
    Эта функция больше не работает, так как get_answer теперь требует user_id и session.
    """
    raise NotImplementedError(
        "generate_response() устарела. Используйте GeminiService.get_answer(prompt, user_id, session, context)"
    )
